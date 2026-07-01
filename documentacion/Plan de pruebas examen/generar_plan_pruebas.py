from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

OUTPUT_DIR = Path(__file__).resolve().parent

HEADERS = [
    "ID",
    "Módulo",
    "Requisito / HU",
    "Tipo de prueba",
    "Acción / procedimiento",
    "Datos de prueba",
    "Resultado esperado",
    "Resultado obtenido",
    "Estado",
    "Evidencia / referencia",
]

ROWS = [
    [
        "TP-01",
        "Backend — Seguridad",
        "Política de contraseña (registro / cambio)",
        "Unitaria automatizada",
        "Ejecutar mvn test; clase PasswordRulesTest (7 casos: válida, corta, sin mayúscula, sin minúscula, sin dígito, sin símbolo, null).",
        "Contraseñas de prueba según cada caso en PasswordRulesTest.java",
        "Solo se aceptan contraseñas que cumplen la política; el resto se rechaza.",
        "Build Success. Captura compartida en §6.1 (EV-UNIT; TP-01 a TP-06).",
        "OK",
        "§6.1 — EV-UNIT: mvn test BUILD SUCCESS (PasswordRulesTest)",
    ],
    [
        "TP-02",
        "Backend — Seguridad",
        "JWT sesión y recuperación",
        "Unitaria automatizada",
        "Ejecutar mvn test; clase JwtServiceTest (claims role/email; token recuperación con purpose).",
        "Usuario mock con rol RECRUITER y email de prueba",
        "Token incluye claims esperados; token recuperación incluye purpose.",
        "Build Success. Captura compartida en §6.1 (EV-UNIT; TP-01 a TP-06).",
        "OK",
        "§6.1 — EV-UNIT: mvn test BUILD SUCCESS (JwtServiceTest)",
    ],
    [
        "TP-03",
        "Backend — Seguridad",
        "Login y bloqueo tras intentos fallidos",
        "Unitaria automatizada",
        "Ejecutar mvn test; clase AuthServiceTest (login válido; bloqueo tras máximo de intentos).",
        "Credenciales válidas e inválidas simuladas con mocks",
        "Login feliz devuelve token; tras 4 intentos fallidos la cuenta queda bloqueada.",
        "Build Success. Captura compartida en §6.1 (EV-UNIT; TP-01 a TP-06).",
        "OK",
        "§6.1 — EV-UNIT: mvn test BUILD SUCCESS (AuthServiceTest)",
    ],
    [
        "TP-04",
        "Backend — CV / Tika",
        "Extracción texto PDF (HU-10)",
        "Unitaria automatizada",
        "Ejecutar mvn test; clase CvTextExtractionServiceTest (PDF con texto seleccionable; PDF sin texto útil).",
        "PDF de prueba con texto y PDF vacío/imagen",
        "Texto normalizado persistido en caso válido; rechazo cuando no hay texto útil.",
        "Build Success. Captura compartida en §6.1 (EV-UNIT; TP-01 a TP-06).",
        "OK",
        "§6.1 — EV-UNIT: mvn test BUILD SUCCESS (CvTextExtractionServiceTest)",
    ],
    [
        "TP-05",
        "Backend — Matching",
        "Score cinco dimensiones (HU-07 / HU-08)",
        "Unitaria automatizada",
        "Ejecutar mvn test; clase ProductMatchServiceTest (promedio 5 dims, modality/workload/experience, sin CV, sin texto, IA no disponible).",
        "Mocks de CV, preferencias candidato, respuesta IA",
        "final_score = promedio de description, title, modality, workload, experience; null en casos sin CV/texto/IA.",
        "Build Success. Captura compartida en §6.1 (EV-UNIT; TP-01 a TP-06).",
        "OK",
        "§6.1 — EV-UNIT: mvn test BUILD SUCCESS (ProductMatchServiceTest)",
    ],
    [
        "TP-06",
        "Backend — Matching",
        "Dimensión años de experiencia",
        "Unitaria automatizada",
        "Ejecutar mvn test; clase ExperienceScoreUtilTest (cumple requisito, escalas por ratio, job sin años).",
        "Años candidato vs años requeridos en oferta",
        "Puntuación coherente con reglas de experiencia del contrato.",
        "Build Success. Captura compartida en §6.1 (EV-UNIT; TP-01 a TP-06).",
        "OK",
        "§6.1 — EV-UNIT: mvn test BUILD SUCCESS (ExperienceScoreUtilTest)",
    ],
    [
        "TP-07",
        "Backend — Seguridad",
        "Auth HTTP y control de roles",
        "Integración automatizada",
        "Ejecutar mvn test -Pintegration; AuthIntegrationTest (login JWT con rol; RECRUITER sin /admin/**; bloqueo tras 4 intentos).",
        "MySQL Testcontainers; usuario RECRUITER y ADMIN sembrados",
        "200 + JWT con rol en login; 403 en admin para RECRUITER; cuenta bloqueada tras intentos fallidos.",
        "Build Success. Captura compartida en §6.1 (EV-INT; TP-07 a TP-10).",
        "OK",
        "§6.1 — EV-INT: mvn test -Pintegration BUILD SUCCESS (AuthIntegrationTest)",
    ],
    [
        "TP-08",
        "Backend + IA",
        "Detalle oferta con score e integración IA",
        "Integración automatizada",
        "Ejecutar mvn test -Pintegration; MatchScoreIntegrationTest (5 dimensiones + final; oferta alineada > no alineada).",
        "Candidato con CV; dos ofertas alineada y desalineada; FastAPI en Testcontainers o WORKSI_IT_AI_URL",
        "Respuesta incluye cinco dimensiones y final_score consistente; oferta alineada puntúa mayor.",
        "Build Success. Captura compartida en §6.1 (EV-INT; TP-07 a TP-10).",
        "OK",
        "§6.1 — EV-INT: mvn test -Pintegration BUILD SUCCESS (MatchScoreIntegrationTest)",
    ],
    [
        "TP-09",
        "Backend — Flujo core",
        "Postulación y mensajería mínima (HU-12 / HU-29–31)",
        "Integración automatizada",
        "Ejecutar mvn test -Pintegration; CoreApiIntegrationTest (postular; listar postulaciones empresa con score; crear conversación y mensajes).",
        "Candidato y reclutador autenticados; oferta publicada",
        "Postulación creada; listado empresa muestra score; conversación y mensajes intercambiados.",
        "Build Success. Captura compartida en §6.1 (EV-INT; TP-07 a TP-10).",
        "OK",
        "§6.1 — EV-INT: mvn test -Pintegration BUILD SUCCESS (CoreApiIntegrationTest)",
    ],
    [
        "TP-10",
        "Backend — Matching",
        "Invariante: skills no alteran score",
        "Integración automatizada",
        "Ejecutar mvn test -Pintegration; MatchInvariantIntegrationTest (cambiar solo skills perfil; cambiar solo job_skills).",
        "Mismo CV y mismos textos IA; mutación solo en skills BD",
        "final_score idéntico antes y después del cambio de skills.",
        "Build Success. Captura compartida en §6.1 (EV-INT; TP-07 a TP-10).",
        "OK",
        "§6.1 — EV-INT: mvn test -Pintegration BUILD SUCCESS (MatchInvariantIntegrationTest)",
    ],
    [
        "TP-11",
        "IA — FastAPI",
        "Contrato POST /match y /health",
        "Unitaria automatizada (pytest)",
        "En producto/ai-service: pytest test_main_fast.py (health; payload corto; score+explanation; CV sin texto tras fairness).",
        "Payloads JSON mínimos y mocks de modelos",
        "/health 200; /match rechaza payload inválido; respuesta con score y explanation en rango; rechazo CV sin texto útil.",
        "Build Success. Captura compartida en §6.1 (EV-IA; TP-11 a TP-13).",
        "OK",
        "§6.1 — EV-IA: pytest passed (test_main_fast.py)",
    ],
    [
        "TP-12",
        "IA — Fairness",
        "Limpieza texto oferta (fairness job)",
        "Unitaria automatizada (pytest)",
        "En producto/ai-service: pytest test_fairness.py (elimina email/RUT; preserva semántica laboral).",
        "Textos de oferta con datos sensibles y contenido laboral",
        "Email y RUT removidos; descripción laboral preservada.",
        "Build Success. Captura compartida en §6.1 (EV-IA; TP-11 a TP-13).",
        "OK",
        "§6.1 — EV-IA: pytest passed (test_fairness.py)",
    ],
    [
        "TP-13",
        "IA — Semántica",
        "Fairness CV, embeddings y explicación",
        "Unitaria automatizada (pytest slow)",
        "En producto/ai-service: pytest -m slow test_semantic_match.py (redacción CV; score en rango; par alineado > no relacionado; explanation no vacía).",
        "Pares CV–oferta alineados y no relacionados; modelos spaCy + Sentence Transformers",
        "Datos sensibles redactados; scores en [0,1]; par alineado mayor; explicación no vacía.",
        "Build Success. Captura compartida en §6.1 (EV-IA; TP-11 a TP-13).",
        "OK",
        "§6.1 — EV-IA: pytest passed (test_semantic_match.py)",
    ],
    [
        "TP-14",
        "Cloud — Admin",
        "Alta empresa y reclutadores (Sprint 2)",
        "Manual funcional (cloud)",
        "En portal web Vercel: login ADMIN; crear empresa multipart; crear reclutador asociado (API vía proxy /api).",
        "Credenciales ADMIN semilla; datos empresa y reclutador de demo",
        "Empresa y reclutador visibles; reclutador puede iniciar sesión.",
        "Login ADMIN, alta empresa y reclutador exitosos; login reclutador verificado. Ver §6.2 (EV-14).",
        "OK",
        "§6.2 — EV-14: capturas web ADMIN",
    ],
    [
        "TP-15",
        "Cloud — Candidato móvil",
        "Registro candidato multipart (HU-01 / HU-21)",
        "Manual funcional (cloud)",
        "App Android con BASE_URL backend Railway: consentimiento HU-21; registro con CV PDF y datos perfil.",
        "CV PDF con texto; email único; skills perfil",
        "Cuenta CANDIDATE operativa; CV procesado; puede autenticarse.",
        "Consentimiento HU-21, registro multipart con CV y perfil completados; candidato operativo. Ver §6.2 (EV-15).",
        "OK",
        "§6.2 — EV-15: capturas móvil registro",
    ],
    [
        "TP-16",
        "Cloud — Reclutador web",
        "Publicación oferta (HU-15)",
        "Manual funcional (cloud)",
        "Portal RECRUITER en Vercel: crear oferta con campos obligatorios y skills.",
        "Sesión reclutador de TP-14",
        "Oferta publicada y visible en listados empresa.",
        "Oferta creada y visible en listado empresa del reclutador. Ver §6.2 (EV-16).",
        "OK",
        "§6.2 — EV-16: capturas publicación oferta",
    ],
    [
        "TP-17",
        "Cloud — Candidato móvil",
        "Feed swipe y postulación (HU-11 / HU-12)",
        "Manual funcional (cloud)",
        "Candidato autenticado contra API Railway: feed con score; swipe APPLY; confirmar postulación.",
        "Candidato TP-15; oferta TP-16",
        "Postulación registrada; estado coherente en app.",
        "Feed con score, swipe APPLY y postulación confirmada. Ver §6.2 (EV-17).",
        "OK",
        "§6.2 — EV-17: capturas feed y postulación",
    ],
    [
        "TP-18",
        "Cloud — Reclutador web",
        "Postulaciones y Detalle de Score (HU-19 / HU-28)",
        "Manual funcional (cloud)",
        "Reclutador en Vercel: listado postulantes; abrir detalle; ver cinco dimensiones y score.",
        "Postulación de TP-17",
        "Postulante visible con score; detalle muestra dimensiones del contrato.",
        "Listado postulantes y Detalle de Score con cinco dimensiones y porcentaje final. Ver §6.2 (EV-18).",
        "OK",
        "§6.2 — EV-18: capturas postulantes y score",
    ],
    [
        "TP-19",
        "Cloud — Match y mensajería",
        "Establecer match y mensajes (HU-29–HU-31)",
        "Manual funcional (cloud)",
        "Reclutador web: Establecer Match con primer mensaje; bandeja Matchs; hilo mensajes. Candidato móvil: aviso post-login; Matchs; responder.",
        "Postulación de TP-17; primer mensaje max 80 caracteres",
        "Conversación creada; mensajes bidireccionales visibles solo para participantes.",
        "Match establecido, mensajería bidireccional web reclutador y móvil candidato verificada. Ver §6.2 (EV-19).",
        "OK",
        "§6.2 — EV-19: capturas match y mensajería",
    ],
    [
        "TP-20",
        "Cloud — Matching",
        "Invariante scoring (evidencia documental)",
        "Manual / documental",
        "En entorno cloud: misma oferta y mismo CV; cambiar solo skills perfil u oferta; comparar score antes/después (complementa TP-10 / Sprint 12).",
        "Dos capturas o logs de consultas con mismos textos IA",
        "Porcentaje final no cambia al mutar solo skills en BD.",
        "Score final idéntico antes y después de cambiar solo skills de perfil u oferta (mismo CV). Ver §6.2 (EV-20).",
        "OK",
        "§6.2 — EV-20: capturas score antes/después",
    ],
    [
        "TP-21",
        "Cloud — Infraestructura",
        "Despliegue Vercel + Railway (Sprint 13)",
        "Manual operativa",
        "Tras deploy: GET /health backend e IA en Railway; portal web Vercel accesible; rewrite /api hacia backend; volumen CV montado en backend.",
        "Servicios Railway (MySQL, backend, ai-service) y proyecto Vercel frontWeb",
        "Health UP; web carga; API responde vía proxy; backend persiste CV.",
        "Health backend e IA OK; web Vercel operativa; proxy /api funcional. Ver §6.2 (EV-21).",
        "OK",
        "§6.2 — EV-21: capturas health y portal web",
    ],
    [
        "TP-22",
        "Cloud — Seguridad",
        "Rutas privadas y JWT",
        "Manual / exploratoria",
        "Contra URL pública del backend Railway: sin token GET rutas protegidas → 401; con JWT rol incorrecto → 403 donde aplique.",
        "Postman o cliente HTTP; tokens ADMIN, RECRUITER, CANDIDATE",
        "Rutas privadas rechazan acceso no autorizado.",
        "GET /api/v1/admin/companies sin token: acceso denegado. Con JWT RECRUITER: 403 Forbidden. Ver §6.2 (EV-22).",
        "OK",
        "§6.2 — EV-22: capturas Postman",
    ],
]

AUTOMATED_EVIDENCE_ANNEX = [
    (
        "EV-UNIT — Unitarios backend (TP-01 a TP-06)",
        "Comando: cd producto/backend → mvn test. Una sola captura BUILD SUCCESS cubre: "
        "TP-01 PasswordRulesTest · TP-02 JwtServiceTest · TP-03 AuthServiceTest · "
        "TP-04 CvTextExtractionServiceTest · TP-05 ProductMatchServiceTest · TP-06 ExperienceScoreUtilTest.",
    ),
    (
        "EV-INT — Integración backend + IA (TP-07 a TP-10)",
        "Comando: cd producto/backend → mvn test -Pintegration (Docker en ejecución). Una sola captura BUILD SUCCESS cubre: "
        "TP-07 AuthIntegrationTest · TP-08 MatchScoreIntegrationTest · TP-09 CoreApiIntegrationTest · "
        "TP-10 MatchInvariantIntegrationTest.",
    ),
    (
        "EV-IA — Servicio IA ai-service (TP-11 a TP-13)",
        "Comando: cd producto/ai-service → pytest (o pytest -m \"not slow\" + pytest -m slow). Una sola captura con todos passed cubre: "
        "TP-11 test_main_fast.py · TP-12 test_fairness.py · TP-13 test_semantic_match.py.",
    ),
]

E2E_EVIDENCE_ANNEX = [
    ("TP-14 / EV-14", "Admin: login, alta empresa, alta reclutador, login reclutador."),
    ("TP-15 / EV-15", "Candidato móvil: consentimiento HU-21, registro con CV, sesión operativa."),
    ("TP-16 / EV-16", "Reclutador: publicación de oferta y listado."),
    ("TP-17 / EV-17", "Candidato móvil: feed con score, swipe y postulación."),
    ("TP-18 / EV-18", "Reclutador: listado postulantes y Detalle de Score."),
    ("TP-19 / EV-19", "Match y mensajería bidireccional (web reclutador y móvil candidato)."),
    ("TP-20 / EV-20", "Invariante scoring: mismo porcentaje tras cambiar solo skills."),
    ("TP-21 / EV-21", "Railway + Vercel: health backend e IA, portal web y proxy /api operativos."),
    ("TP-22 / EV-22", "Postman: acceso denegado sin token y 403 con JWT RECRUITER en /admin."),
]

INTRO_PARAGRAPHS = [
    "Proyecto: WorkSí — plataforma de reclutamiento (Spring Boot, FastAPI, MySQL, React, Android).",
    "Este documento constituye el plan de pruebas de software alineado a la Evaluación Final Transversal TPY1101 y al cierre Sprint 12–13 del proyecto.",
    "Alcance: pruebas automatizadas de backend e IA (Sprint 12) y validación manual del despliegue cloud y demo del flujo core (Sprint 13). Estado del plan: 22/22 casos ejecutados.",
    "Entorno de pruebas: local con Java 21, Maven, Python 3.11, Docker Desktop (Sprint 12). Integración backend usa Testcontainers (MySQL + imagen ai-service) o WORKSI_IT_AI_URL apuntando a IA en localhost:8000. Sprint 13: entorno cloud Vercel (web) + Railway (MySQL, backend, IA); app móvil con BASE_URL HTTPS del backend.",
    "Comandos de referencia (producto/README.md §13): mvn test (unitarios backend); mvn test -Pintegration (integración); pytest o pytest -m \"not slow\" (IA).",
    "Criterio de éxito general: cada caso de la tabla cumple su resultado esperado; bugs críticos (caídas, 500 en flujo core en cloud, proxy /api roto, BD no migrada, score inconsistente, mensajes a terceros) deben quedar en cero antes de la demo stakeholder.",
    "Evidencias (TP-01 a TP-22): capturas en la sección 6 de este documento. TP-01 a TP-13: tres capturas de terminal (EV-UNIT, EV-INT, EV-IA). TP-14 a TP-22: capturas de validación cloud y demo del producto.",
]

SUMMARY_STATS = [
    ("Pruebas automatizadas backend (unitarias)", "6 filas (TP-01 a TP-06) — 1 captura EV-UNIT en §6.1"),
    ("Pruebas automatizadas backend (integración)", "4 filas (TP-07 a TP-10) — 1 captura EV-INT en §6.1"),
    ("Pruebas automatizadas IA (pytest)", "3 filas (TP-11 a TP-13) — 1 captura EV-IA en §6.1"),
    ("Pruebas manuales cloud y demo (Sprint 13)", "9 filas (TP-14 a TP-22) — evidencias en §6 de este documento (EV-14 a EV-22)"),
    ("Total casos en plan", "22 casos — 22 OK"),
]


def style_excel(ws):
    header_fill = PatternFill("solid", fgColor="1F4E79")
    header_font = Font(color="FFFFFF", bold=True, size=10)
    thin = Side(style="thin", color="CCCCCC")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    widths = [8, 18, 22, 16, 38, 24, 28, 22, 10, 28]

    for col, width in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(col)].width = width

    for col, header in enumerate(HEADERS, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = border

    ws.row_dimensions[1].height = 36
    ws.freeze_panes = "A2"

    ok_fill = PatternFill("solid", fgColor="E2EFDA")
    pending_fill = PatternFill("solid", fgColor="FFF2CC")

    for row_idx, row in enumerate(ROWS, 2):
        for col_idx, value in enumerate(row, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = border
            if col_idx == 9:
                if value == "OK":
                    cell.fill = ok_fill
                elif value == "Pendiente":
                    cell.fill = pending_fill
        ws.row_dimensions[row_idx].height = 72


def build_excel(path: Path):
    wb = Workbook()
    ws_plan = wb.active
    ws_plan.title = "Plan de pruebas"

    ws_plan["A1"] = "WorkSí — Plan de pruebas de software (TPY1101)"
    ws_plan["A1"].font = Font(bold=True, size=14)
    ws_plan.merge_cells("A1:J1")

    intro_row = 3
    for i, text in enumerate(INTRO_PARAGRAPHS):
        ws_plan.cell(row=intro_row + i, column=1, value=text)
        ws_plan.merge_cells(start_row=intro_row + i, start_column=1, end_row=intro_row + i, end_column=10)

    table_start = intro_row + len(INTRO_PARAGRAPHS) + 2
    for col, header in enumerate(HEADERS, 1):
        ws_plan.cell(row=table_start, column=col, value=header)

    header_fill = PatternFill("solid", fgColor="1F4E79")
    header_font = Font(color="FFFFFF", bold=True, size=10)
    thin = Side(style="thin", color="CCCCCC")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    widths = [8, 18, 22, 16, 38, 24, 28, 22, 10, 28]

    for col, width in enumerate(widths, 1):
        ws_plan.column_dimensions[get_column_letter(col)].width = width

    for col in range(1, len(HEADERS) + 1):
        cell = ws_plan.cell(row=table_start, column=col)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = border

    ok_fill = PatternFill("solid", fgColor="E2EFDA")
    pending_fill = PatternFill("solid", fgColor="FFF2CC")

    for i, row in enumerate(ROWS):
        r = table_start + 1 + i
        for col, value in enumerate(row, 1):
            cell = ws_plan.cell(row=r, column=col, value=value)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = border
            if col == 9:
                if value == "OK":
                    cell.fill = ok_fill
                elif value == "Pendiente":
                    cell.fill = pending_fill
        ws_plan.row_dimensions[r].height = 72

    ws_plan.freeze_panes = ws_plan.cell(row=table_start + 1, column=1).coordinate

    ws_resumen = wb.create_sheet("Resumen")
    ws_resumen["A1"] = "Resumen del plan"
    ws_resumen["A1"].font = Font(bold=True, size=12)
    ws_resumen.column_dimensions["A"].width = 42
    ws_resumen.column_dimensions["B"].width = 50
    for i, (label, value) in enumerate(SUMMARY_STATS, 3):
        ws_resumen.cell(row=i, column=1, value=label).font = Font(bold=True)
        ws_resumen.cell(row=i, column=2, value=value)

    ws_cmds = wb.create_sheet("Comandos")
    ws_cmds["A1"] = "Comandos de ejecución"
    ws_cmds["A1"].font = Font(bold=True, size=12)
    ws_cmds.column_dimensions["A"].width = 28
    ws_cmds.column_dimensions["B"].width = 55
    commands = [
        ("Unitarios backend", "cd producto/backend && mvn test"),
        ("Integración backend + IA", "cd producto/backend && mvn test -Pintegration"),
        ("IA con IA ya en localhost:8000", "$env:WORKSI_IT_AI_URL='http://localhost:8000'; mvn test -Pintegration"),
        ("IA pytest (rápidas)", "cd producto/ai-service && pytest -m \"not slow\""),
        ("IA pytest (completas)", "cd producto/ai-service && pytest"),
        ("Docker entorno", "cd producto && docker compose up"),
    ]
    for i, (label, cmd) in enumerate(commands, 3):
        ws_cmds.cell(row=i, column=1, value=label).font = Font(bold=True)
        ws_cmds.cell(row=i, column=2, value=cmd)

    wb.save(path)


def add_table_to_doc(doc, rows, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row in rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)


def build_word(path: Path):
    doc = Document()
    title = doc.add_heading("WorkSí — Plan de pruebas de software", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("1. Introducción", level=1)
    for text in INTRO_PARAGRAPHS:
        doc.add_paragraph(text)

    doc.add_heading("2. Resumen cuantitativo", level=1)
    for label, value in SUMMARY_STATS:
        doc.add_paragraph(f"{label}: {value}", style="List Bullet")

    doc.add_heading("3. Comandos de ejecución", level=1)
    doc.add_paragraph("Unitarios backend: cd producto/backend → mvn test")
    doc.add_paragraph("Integración backend + IA: cd producto/backend → mvn test -Pintegration")
    doc.add_paragraph("IA (rápidas): cd producto/ai-service → pytest -m \"not slow\"")
    doc.add_paragraph("IA (completas): cd producto/ai-service → pytest")
    doc.add_paragraph("Entorno Docker: cd producto → docker compose up")

    doc.add_heading("4. Tabla del plan de pruebas", level=1)
    doc.add_paragraph(
        "La tabla siguiente documenta cada caso con ID, módulo, requisito, tipo, procedimiento, "
        "datos, resultado esperado, resultado obtenido, estado y evidencia. "
        "Los casos TP-01 a TP-13 corresponden a pruebas automatizadas (evidencias agrupadas EV-UNIT, EV-INT, EV-IA en §6.1); "
        "TP-14 a TP-22 a la validación manual en cloud y demo del flujo core (§6.2, EV-14 a EV-22)."
    )
    add_table_to_doc(doc, ROWS, HEADERS)

    doc.add_heading("5. Tabla de mejoras (plantilla)", level=1)
    doc.add_paragraph(
        "Completar tras la validación cloud (Sprint 13) si se detectan hallazgos (requisito ítem 16 pauta TPY1101)."
    )
    mejoras_headers = [
        "ID hallazgo",
        "Caso TP",
        "Descripción del problema",
        "Severidad",
        "Acción correctiva",
        "Estado re-test",
        "Evidencia",
    ]
    mejoras_rows = [
        ["H-01", "—", "Sin hallazgos críticos en la validación cloud (TP-14 a TP-22)", "—", "—", "N/A", "—"],
    ]
    add_table_to_doc(doc, mejoras_rows, mejoras_headers)

    doc.add_heading("6. Anexo de evidencias (TP-01 a TP-22)", level=1)
    doc.add_paragraph(
        "Capturas de la ejecución de pruebas. TP-01 a TP-13: tres capturas de terminal agrupadas por tipo de suite (§6.1). "
        "TP-14 a TP-22: capturas de validación cloud y demo del producto (§6.2). La tabla del §4 mantiene referencia EV-01 a EV-22 por caso; "
        "en §6.1 varios TP comparten la misma imagen según el comando ejecutado."
    )
    doc.add_heading("6.1 Pruebas automatizadas — terminal (TP-01 a TP-13)", level=2)
    for title, description in AUTOMATED_EVIDENCE_ANNEX:
        doc.add_heading(title, level=3)
        doc.add_paragraph(description)
        doc.add_paragraph("[Insertar una captura de terminal BUILD SUCCESS / passed aquí]")
    doc.add_heading("6.2 Validación cloud y demo — producto (Sprint 13, TP-14 a TP-22)", level=2)
    for title, description in E2E_EVIDENCE_ANNEX:
        doc.add_heading(title, level=3)
        doc.add_paragraph(description)
        doc.add_paragraph("[Insertar capturas de pantalla aquí]")

    doc.add_heading("7. Conclusiones", level=1)
    doc.add_paragraph(
        "El plan de pruebas cubre 22 casos en estado OK: suite automatizada backend/IA (TP-01 a TP-13) "
        "y validación manual del despliegue cloud Vercel + Railway con demo del flujo core admin–reclutador–candidato–match–mensajería, "
        "health de servicios y control de acceso JWT (TP-14 a TP-22). Evidencias automatizadas agrupadas en §6.1 (EV-UNIT, EV-INT, EV-IA); cloud y demo en §6.2 (EV-14 a EV-22)."
    )

    for section in doc.sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    doc.save(path)


def main():
    xlsx_path = OUTPUT_DIR / "plan-de-pruebas.xlsx"
    docx_path = OUTPUT_DIR / "plan-de-pruebas.docx"
    build_excel(xlsx_path)
    build_word(docx_path)
    print(f"Generado: {xlsx_path}")
    print(f"Generado: {docx_path}")


if __name__ == "__main__":
    main()
